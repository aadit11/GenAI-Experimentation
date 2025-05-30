from langchain.chat_models import AzureChatOpenAI
from langchain.schema import HumanMessage
import os
from datetime import datetime
from docx import Document

openai_api_key = ""

def process_text_for_safety(text):
    safe_text = process_text_for_safety(text)
    print(safe_text, end=' ', flush=True)

llm = AzureChatOpenAI(
    deployment_name="",
    openai_api_version="",
    openai_api_key=openai_api_key,
    openai_api_base="",
)

print(datetime.now())

doc_path = "data\file.docx"

def read_word_document(doc_path):
    doc = Document(doc_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def extract_tables_from_document(doc_path):
    doc = Document(doc_path)
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        tables.append(table_data)
    return tables

def handle_input_from_terminal_and_document(doc_path, previous_query=None):
    document_text = read_word_document(doc_path)
    
    tables = extract_tables_from_document(doc_path)
    
    user_query = input("Enter your query (or press Enter to use the previous query):\n") or previous_query
    
    previous_query = user_query
    
    user_message = HumanMessage(content=f"{document_text}\n{user_query}", tables=tables)
    
    response_generator = llm.stream([user_message])

    for chunk in response_generator:
        print(chunk.content, end='', flush=True)
    
    return previous_query

print(datetime.now())

document_path = "data\file.docx"
previous_query = None

while True:
    previous_query = handle_input_from_terminal_and_document(document_path, previous_query)
    print(datetime.now())
