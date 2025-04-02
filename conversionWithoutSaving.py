from langchain.chat_models import AzureChatOpenAI
from langchain.schema import HumanMessage
from datetime import datetime
from docx import Document
from bs4 import BeautifulSoup
import tabula
import fitz
import pandas as pd

openai_api_key = ""

llm = AzureChatOpenAI(
    deployment_name="",
    openai_api_version="",
    openai_api_key=openai_api_key,
    openai_api_base="",
)

def convert_pdf_to_dataframe(pdf_path):
    tables = tabula.read_pdf(pdf_path, pages='all', multiple_tables=True)
    df = pd.concat(tables, ignore_index=True)
    return df

def extract_text_from_pdf(pdf_path):
    doc = fitz.open(pdf_path)
    text = ""
    for page_num in range(doc.page_count):
        page = doc[page_num]
        text += page.get_text("text")
    return text

def handle_input_from_terminal(pdf_path, previous_context=None):
    df = convert_pdf_to_dataframe(pdf_path)
    user_query = input("Enter your query (or press Enter to skip):\n")
    csv_content = df.to_csv(index=False)

    # Load CSV content into a Document
    document = Document()
    document.add_paragraph(csv_content)

    split_text = "\n".join([paragraph.text for paragraph in document.paragraphs])

    user_message = HumanMessage(content=split_text, context=previous_context)

    if user_query:
        user_message.content += f"\nUser Query: {user_query}"

    response_generator = llm.stream([user_message])

    for chunk in response_generator:
        print(chunk.content, end='', flush=True)

    return user_message.context

pdf_path = "data\file.pdf"
previous_context = None

while True:
    previous_context = handle_input_from_terminal(pdf_path, previous_context)
    print(datetime.now())